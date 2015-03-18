# -*- coding: utf-8 -*-
"""
Created on Wed Mar 18 06:26:25 2015

@author: Alex
"""
plt.ioff()
plt.close('all')
from docx import *
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
##  Create Document
document = Document()

######## SOME TOOLS
def add_figure_caption(fig_num=str(len(document.inline_shapes)),caption=''):
    cap = document.add_paragraph("Figure "+fig_num+". "+caption)
    cap.paragraph_style = 'caption'
    return
    
def dataframe_to_table(df=pd.DataFrame(),table_num=str(len(document.tables)+1),caption='',fontsize=11):
    table = document.add_table(rows=1, cols=len(df.columns)) 
    ## Merge all cells in top row and add caption text
    table_caption = table.rows[0].cells[0].merge(table.rows[0].cells[len(df.columns)-1])
    table_caption.text = "Table "+table_num+". "+caption
    ## Add  header
    header_row = table.add_row().cells
    col_count =0 ## counter  to iterate over columns
    for col in  header_row:
        col.text = df.columns[col_count] #df.columns[0] is the index
        col_count+=1
    ## Add data by  iterating over the DataFrame rows, then using a dictionary of DataFrame column labels to extract data
    col_labels = dict(zip(range(len(df.columns)),df.columns.tolist())) ## create dictionary where '1  to  n' is key for DataFrame columns
    for row in df.iterrows():  ## iterate over  rows in  DataFrame
        #print row[1]
        row_cells = table.add_row().cells ## Add a row to the  table
        col_count  = 0
        for cell in row_cells: ## iterate over the columns in the row
            #print row[1][str(col_labels[col_count])]
            cell.text = str(row[1][str(col_labels[col_count])]) ## and plug in data using a dictionary to  get column labels for DataFrame
            col_count+=1
    ## Format Table Style  
    table.style = 'TableGrid' 
    table.style.font.size = Pt(fontsize)
    table.autofit
    #table.num = str(len(document.tables)+1)
    return table
    
def add_equation(eq_table):
    t = document.add_table(rows=len(eq_table.rows),cols=len(eq_table.columns))
    t.rows[0].cells[1].text = eq_table.rows[0].cells[1].text
    t.rows[0].cells[2].text = eq_table.rows[0].cells[2].text   
    if len(eq_table.rows)>1:
        t.rows[1].cells[0].merge(t.rows[1].cells[2]).text = eq_table.rows[1].cells[0].text
    t.style = 'TableGrid'   
    t.autofit
    return t
###################################################################################################################################################################    


#### TABLES ########################################################################################################################################################
### Landcover_Table
table_count=0
def tab_count():
    global table_count
    table_count+=1
    return str(table_count)
# Prepare LULC Data
table_one = table_one_function()
table_one.table_num = str(tab_count())    


#### FIGURES ########################################################################################################################################################
figure_count=0
def fig_count():
    global figure_count
    figure_count+=1
    return str(figure_count)
## INTRODUCTION
#### Study Area Map
Study_Area_map = {'filename':maindir+'Figures/Maps/Study_Area.tif', 'fig_num':str(fig_count())}


###### EQUATIONS ############################################################################################################################################
equation_count=0
def eq_count():
    global equation_count
    equation_count+=1
    return str(equation_count)
    
Equations = Document(maindir+'/Manuscript/Equations.docx').tables

## SSYev = Q*SSC 
Equation_one = Equations[0].table
Equation_one.eq_num = eq_count()

##### APPENDIX #################################################################################################################################
#### Appendix
table_count,figure_count,equation_count=0, 0, 0
### Appendix Table One
Appendix_table_one = table_function() ## function to create table data
Appendix_table_one.table_num =str(tab_count())

############################################################################################################################################

#### TITLE
title_title = document.add_heading('TITLE:',level=1)
title_title.paragraph_format.space_before = 0
title = document.add_heading('Contributions of human activities to suspended sediment yield during storm events from a steep, small, tropical watershed',level=1)
title.paragraph_format.space_before = 0
## subscript/superscript words
document.add_paragraph("")

#### ABSTRACT
abstract_title = document.add_heading('ABSTRACT',level=2)
abstract_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
abstract = document.add_paragraph('Abstract text goes here....')


#### INTRODUCTION
introduction_title = document.add_heading('Introduction',level=2)
document.add_paragraph("Intro..")

#### STUDY AREA
study_area_title = document.add_heading('Study Area',level=2)
## Study Area map
if 'Study_Area_map' in locals():
    document.add_picture(Study_Area_map['filename'],width=Inches(6))
    add_figure_caption(Study_Area_map['fig_num'],"figure caption.")
    
#### METHODS
methods_title = document.add_heading('Methods',level=2)
document.add_paragraph("Methods...")


#### RESULTS ####
results_title = document.add_heading('Results and Discussion',level=2)
## Field data collection
document.add_heading('Field Data Collection',level=3)

#### CONCLUSION
conclusion_title=document.add_heading('Conclusion',level=2)
conclusion_text = document.add_paragraph("In conclusion...")

#### Appendix
document.add_page_break()
document.add_heading('APPENDIX',level=2)
## Appendix stuff goes here...

## Save Document
document.save(maindir+'Manuscript/DRAFT.docx')

## Clean up any open figures
plt.close('all')

